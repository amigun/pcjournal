from datetime import datetime, timedelta
from docx import Document

from .models import Files

def generate(date, schedule):
    WEEKDAYS = {
        0: 'Понедельник',
        1: 'Вторник',
        2: 'Среда',
        3: 'Четверг',
        4: 'Пятница',
        5: 'Суббота',
    }

    document = Document()

    print(schedule)

    for day in schedule:
        weekday = WEEKDAYS[day[0].weekday()]
        document.add_heading(f'{day[0].strftime("%d %b")}, {weekday}')

        for task in day[1]:
            document.add_heading(str(task[0]), 3)
            document.add_heading(str(f'\t{task[1]}:'), 4)
            document.add_paragraph(str(f'\t\t{task[2]}'))
            document.add_heading(str(f'\t\t{task[3]}'), 4)

    document.save(f'files/{date}.docx')

    file = Files(title=date, path=f'files/{date}.docx')
    file.save()
